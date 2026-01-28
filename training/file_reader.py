"""
Universal File Reader

Supports reading multiple file formats:
- Documents: PDF, DOCX, TXT, MD
- Code: PY, CPP, CU, HPP, H, C, JAVA, JS, TS
- Data: JSON, CSV, JSONL, XML
- Other text formats

Extracts text content for training the cognitive field.
"""
try: import usage_tracker; usage_tracker.track(__file__)
except: pass

import os
from pathlib import Path
from typing import Optional, List
import json

# Optional dependencies
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    try:
        import pdfplumber
        PDF_AVAILABLE = True
        PyPDF2 = None
    except ImportError:
        PDF_AVAILABLE = False
        PyPDF2 = None
        print("Warning: PDF support not available. Install: pip install PyPDF2 or pdfplumber")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: DOCX support not available. Install: pip install python-docx")


class UniversalFileReader:
    """
    Reads content from various file formats.

    Supported formats:
    - PDF (.pdf)
    - Word (.docx)
    - Markdown (.md)
    - Code (.py, .cpp, .cu, .hpp, .h, .c, .java, .js, .ts, .rs, .go)
    - Data (.json, .jsonl, .csv, .xml)
    - Text (.txt)
    """

    def __init__(self):
        self.supported_extensions = {
            # Documents
            '.pdf': self._read_pdf,
            '.docx': self._read_docx,
            '.doc': self._read_docx,
            '.txt': self._read_text,
            '.md': self._read_text,
            '.markdown': self._read_text,
            '.rst': self._read_text,

            # Code files
            '.py': self._read_text,
            '.cpp': self._read_text,
            '.cu': self._read_text,
            '.hpp': self._read_text,
            '.h': self._read_text,
            '.c': self._read_text,
            '.java': self._read_text,
            '.js': self._read_text,
            '.ts': self._read_text,
            '.rs': self._read_text,
            '.go': self._read_text,
            '.sh': self._read_text,
            '.bat': self._read_text,

            # Data files
            '.json': self._read_json,
            '.jsonl': self._read_jsonl,
            '.csv': self._read_csv,
            '.xml': self._read_text,
            '.yaml': self._read_text,
            '.yml': self._read_text,
            '.toml': self._read_text,
            '.ini': self._read_text,
            '.cfg': self._read_text,

            # Other
            '.log': self._read_text,
            '.html': self._read_text,
            '.htm': self._read_text,
        }

    def can_read(self, file_path: str) -> bool:
        """Check if file format is supported."""
        ext = Path(file_path).suffix.lower()
        return ext in self.supported_extensions

    def read(self, file_path: str) -> str:
        """
        Read file and return text content.

        Args:
            file_path: Path to file

        Returns:
            Extracted text content

        Raises:
            ValueError: If file format not supported
            FileNotFoundError: If file doesn't exist
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()

        if ext not in self.supported_extensions:
            raise ValueError(f"Unsupported file format: {ext}")

        reader_func = self.supported_extensions[ext]

        try:
            content = reader_func(file_path)
            return content
        except Exception as e:
            print(f"Error reading {file_path}: {e}")
            print(f"Attempting fallback text read...")
            try:
                return self._read_text(file_path)
            except:
                raise ValueError(f"Failed to read {file_path}: {e}")

    def _read_text(self, file_path: str) -> str:
        """Read plain text file with multiple encoding attempts."""
        encodings = ['utf-8', 'latin-1', 'ascii', 'cp1252']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
            except Exception as e:
                print(f"Error with encoding {encoding}: {e}")
                continue

        # Last resort: binary read and ignore errors
        with open(file_path, 'rb') as f:
            return f.read().decode('utf-8', errors='ignore')

    def _read_pdf(self, file_path: str) -> str:
        """Read PDF file."""
        if not PDF_AVAILABLE:
            raise ImportError("PDF support not available. Install: pip install PyPDF2 or pdfplumber")

        # Try pdfplumber first (better text extraction)
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            return '\n\n'.join(text_parts)
        except:
            pass

        # Fallback to PyPDF2
        if PyPDF2:
            text_parts = []
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            return '\n\n'.join(text_parts)

        raise ImportError("Could not read PDF with available libraries")

    def _read_docx(self, file_path: str) -> str:
        """Read DOCX file."""
        if not DOCX_AVAILABLE:
            raise ImportError("DOCX support not available. Install: pip install python-docx")

        doc = Document(file_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            text_parts.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)

        return '\n'.join(text_parts)

    def _read_json(self, file_path: str) -> str:
        """Read JSON file and convert to text."""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Convert to readable text
        if isinstance(data, dict):
            return self._dict_to_text(data)
        elif isinstance(data, list):
            return '\n'.join(self._dict_to_text(item) if isinstance(item, dict) else str(item) for item in data)
        else:
            return str(data)

    def _read_jsonl(self, file_path: str) -> str:
        """Read JSONL file (one JSON per line)."""
        text_parts = []
        with open(file_path, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if line:
                    try:
                        data = json.loads(line)
                        if isinstance(data, dict):
                            text_parts.append(self._dict_to_text(data))
                        else:
                            text_parts.append(str(data))
                    except json.JSONDecodeError:
                        text_parts.append(line)
        return '\n'.join(text_parts)

    def _read_csv(self, file_path: str) -> str:
        """Read CSV file."""
        import csv
        text_parts = []

        with open(file_path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            for row in reader:
                text_parts.append(' | '.join(row))

        return '\n'.join(text_parts)

    def _dict_to_text(self, data: dict, prefix: str = '') -> str:
        """Convert dictionary to readable text."""
        text_parts = []

        for key, value in data.items():
            if isinstance(value, dict):
                text_parts.append(f"{prefix}{key}:")
                text_parts.append(self._dict_to_text(value, prefix + '  '))
            elif isinstance(value, list):
                text_parts.append(f"{prefix}{key}: {', '.join(map(str, value))}")
            else:
                text_parts.append(f"{prefix}{key}: {value}")

        return '\n'.join(text_parts)

    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions."""
        return sorted(self.supported_extensions.keys())


def open_file_dialog(title: str = "Select File", filetypes: Optional[List] = None) -> Optional[str]:
    """
    Open GUI file picker dialog with Windows-native appearance.

    Args:
        title: Dialog title
        filetypes: List of (name, pattern) tuples, e.g., [("Text files", "*.txt")]

    Returns:
        Selected file path, or None if cancelled
    """
    try:
        import tkinter as tk
        from tkinter import filedialog
        import sys
    except ImportError:
        print("Error: tkinter not available. Cannot open file dialog.")
        return None

    try:
        # Create hidden root window
        root = tk.Tk()

        # Windows-native appearance and DPI awareness
        if sys.platform == 'win32':
            try:
                from ctypes import windll
                windll.shcore.SetProcessDpiAwareness(1)
            except:
                pass
            root.tk.call('tk', 'scaling', 1.5)

        # Get screen dimensions
        screen_width = root.winfo_screenwidth()
        screen_height = root.winfo_screenheight()

        # Larger dialog, centered
        dialog_width = 1000
        dialog_height = 700
        x = (screen_width - dialog_width) // 2
        y = (screen_height - dialog_height) // 2

        # Set position and bring to front
        root.geometry(f"{dialog_width}x{dialog_height}+{x}+{y}")
        root.withdraw()
        root.attributes('-topmost', True)
        root.lift()
        root.focus_force()

        # Default filetypes
        if filetypes is None:
            filetypes = [
                ("All Supported", "*.pdf *.docx *.txt *.md *.py *.cpp *.cu *.hpp *.json *.csv"),
                ("PDF files", "*.pdf"),
                ("Word documents", "*.docx *.doc"),
                ("Text files", "*.txt *.md"),
                ("Code files", "*.py *.cpp *.cu *.hpp *.h *.c"),
                ("Data files", "*.json *.jsonl *.csv"),
                ("All files", "*.*")
            ]

        # Open dialog
        file_path = filedialog.askopenfilename(
            title=f"{title} - Liorhybrid",
            filetypes=filetypes
        )

        # Ensure proper cleanup
        root.update()
        root.destroy()

        return file_path if file_path else None

    except Exception as e:
        print(f"Error opening file dialog: {e}")
        try:
            root.destroy()
        except:
            pass
        return None


def open_multiple_files_dialog(title: str = "Select Files", filetypes: Optional[List] = None) -> List[str]:
    """
    Open GUI file picker dialog for MULTIPLE file selection with Windows-native appearance.

    Args:
        title: Dialog title
        filetypes: List of (name, pattern) tuples, e.g., [("Text files", "*.txt")]

    Returns:
        List of selected file paths, or empty list if cancelled
    """
    try:
        import tkinter as tk
        from tkinter import filedialog
        import sys
    except ImportError:
        print("ERROR: tkinter not available. Cannot open file dialog.")
        print("This may be a WSL/display issue. Use manual mode instead.")
        return []

    try:
        # Create hidden root window
        root = tk.Tk()

        # Windows-native appearance and DPI awareness
        if sys.platform == 'win32':
            try:
                from ctypes import windll
                windll.shcore.SetProcessDpiAwareness(1)
            except:
                pass
            root.tk.call('tk', 'scaling', 1.5)

        # Get screen dimensions
        screen_width = root.winfo_screenwidth()
        screen_height = root.winfo_screenheight()

        # Larger dialog, centered
        dialog_width = 1000
        dialog_height = 700
        x = (screen_width - dialog_width) // 2
        y = (screen_height - dialog_height) // 2

        # Set position and bring to front
        root.geometry(f"{dialog_width}x{dialog_height}+{x}+{y}")
        root.withdraw()
        root.attributes('-topmost', True)
        root.lift()
        root.focus_force()

        # Default filetypes
        if filetypes is None:
            filetypes = [
                ("All Supported", "*.pdf *.docx *.txt *.md *.py *.cpp *.cu *.hpp *.json *.csv"),
                ("PDF files", "*.pdf"),
                ("Word documents", "*.docx *.doc"),
                ("Text files", "*.txt *.md"),
                ("Code files", "*.py *.cpp *.cu *.hpp *.h *.c"),
                ("Data files", "*.json *.jsonl *.csv"),
                ("All files", "*.*")
            ]

        # Open dialog (note: askopenfilenames - plural!)
        file_paths = filedialog.askopenfilenames(
            title=f"{title} - Liorhybrid",
            filetypes=filetypes
        )

        # Ensure proper cleanup
        root.update()
        root.destroy()

        return list(file_paths) if file_paths else []

    except Exception as e:
        print(f"ERROR opening file dialog: {e}")
        print(f"Exception type: {type(e).__name__}")
        import traceback
        traceback.print_exc()
        try:
            root.destroy()
        except:
            pass
        return []


def read_file_with_dialog(title: str = "Select Training Data") -> Optional[str]:
    """
    Open file dialog and read selected file.

    Returns:
        File content as text, or None if cancelled/error
    """
    file_path = open_file_dialog(title)

    if not file_path:
        return None

    reader = UniversalFileReader()

    try:
        print(f"\nReading file: {file_path}")
        content = reader.read(file_path)
        print(f"✓ Successfully read {len(content)} characters")
        return content, file_path
    except Exception as e:
        print(f"✗ Error reading file: {e}")
        return None


def demo_file_reader():
    """Demonstrate file reader capabilities."""
    print("\n" + "="*80)
    print("UNIVERSAL FILE READER - Demo")
    print("="*80)

    reader = UniversalFileReader()

    print("\nSupported file formats:")
    for ext in reader.get_supported_extensions():
        print(f"  {ext}")

    print("\n" + "="*80)
    print("Opening file picker...")
    print("="*80)

    result = read_file_with_dialog("Select a file to read")

    if result:
        content, file_path = result
        print(f"\nFile: {file_path}")
        print(f"Size: {len(content)} characters")
        print(f"\nFirst 500 characters:")
        print("-" * 80)
        print(content[:500])
        print("-" * 80)
    else:
        print("\nNo file selected or error occurred.")


if __name__ == "__main__":
    demo_file_reader()


# =============================================================================
# DOCX READER (no lxml dependency - uses stdlib zipfile + xml.etree)
# =============================================================================

def read_docx_simple(file_path: str) -> str:
    """
    Read text from a DOCX file using only stdlib (no python-docx/lxml needed).
    
    DOCX files are ZIP archives containing XML. This extracts the main document text.
    """
    import zipfile
    import xml.etree.ElementTree as ET
    
    WORD_NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    
    with zipfile.ZipFile(file_path, 'r') as docx:
        # Main document content is in word/document.xml
        xml_content = docx.read('word/document.xml')
        tree = ET.fromstring(xml_content)
        
        # Extract all text from <w:t> elements
        paragraphs = []
        for paragraph in tree.iter(f'{WORD_NAMESPACE}p'):
            texts = [node.text for node in paragraph.iter(f'{WORD_NAMESPACE}t') if node.text]
            if texts:
                paragraphs.append(''.join(texts))
        
        return '\n'.join(paragraphs)
